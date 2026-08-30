from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT = Path(__file__).resolve().parent
OUT = ROOT / 'استبيان تحليل متطلبات نظام المراسلات والأرشيف - بلدية المغازي.docx'
LOGO = ROOT / 'assets' / 'images' / 'maghazi-logo.jpeg'

MAROON = '8B1E1E'; GOLD = 'B88A2D'; NAVY = '24364A'; BLUE = '2E74B5'
MUTED = '667085'; LIGHT = 'F5F7F9'; PALE = 'FDF2F2'; LINE = 'D9DEE5'; WHITE = 'FFFFFF'

sections = [
('1. الهدف والنطاق والنجاح', [
'ما المشكلة الأساسية التي تريد البلدية حلّها بالنظام؟ وما أكثر شيء يسبب تأخيرًا أو ضياعًا اليوم؟',
'هل النظام سيغطي المراسلات والأرشيف فقط، أم يشمل مهام الموظفين والتعاميم والقرارات ومحاضر الاجتماعات أيضًا؟',
'من هم المستخدمون المستهدفون في المرحلة الأولى؟ وهل ستشارك جهات أو مكاتب خارج مبنى البلدية؟',
'ما النتائج التي ستعتبرون عندها المشروع ناجحًا بعد ثلاثة أشهر من التشغيل؟',
'ما الوظائف التي لا يمكن إطلاق النظام بدونها؟ وما الوظائف التي يمكن تأجيلها لمرحلة ثانية؟',
'هل المطلوب إلغاء الورق بالكامل أم تشغيل إلكتروني وورقي بالتوازي؟ وكم تستمر الفترة الانتقالية؟',
'من صاحب القرار النهائي في اعتماد المتطلبات والتغييرات وقبول النظام؟'
]),
('2. الواقع الحالي وحجم العمل', [
'اشرحوا دورة كتاب وارد حقيقية من لحظة استلامه حتى إغلاقه وأرشفته، مع ذكر الأشخاص والخطوات.',
'اشرحوا دورة كتاب صادر حقيقية من إعداد المسودة حتى الاعتماد والتسليم والأرشفة.',
'كيف تتم المراسلات الداخلية حاليًا؟ وما الفرق عندكم بين مذكرة وتعميم وتكليف وإحالة؟',
'كم متوسط الكتب الواردة والصادرة والمراسلات الداخلية يوميًا وشهريًا؟ وما مواسم الذروة؟',
'ما السجلات أو ملفات Excel أو الدفاتر المستخدمة حاليًا؟ ومن يحتفظ بها؟',
'ما الأخطاء المتكررة حاليًا: تكرار أرقام، فقد مرفقات، تأخر، تحويل خاطئ، أو صعوبة بحث؟',
'هل توجد معاملات قديمة يجب استيرادها؟ ما عدد السنوات وحجم الملفات وصيغها وجودة الفهرسة؟',
'ما الحالات الاستثنائية الحالية التي لا تتبع المسار المعتاد؟'
]),
('3. الهيكل الإداري والمستخدمون', [
'زوّدونا بالقائمة الرسمية للأقسام والدوائر والوحدات ومسؤوليها، وهل يوجد هيكل فرعي داخل كل قسم؟',
'ما مصدر بيانات الموظفين: نظام الموقع الأساسي، الموارد البشرية، أم إدخال يدوي؟',
'هل يمكن أن ينتمي الموظف لأكثر من قسم أو يعمل بالإنابة عن موظف آخر؟',
'ما الأدوار التشغيلية المطلوبة: ديوان، موظف، رئيس قسم، مدير بلدية، أرشيف، مدقق، مراقب؟',
'من يستطيع الاطلاع على مراسلات القسم كاملة، ومن يرى فقط ما تم تحويله إليه؟',
'كيف تُعالج إجازة الموظف أو نقله أو انتهاء خدمته؟ ومن يستلم المراسلات المفتوحة؟',
'هل توجد مجموعات تحويل ثابتة مثل لجنة الطوارئ أو لجنة المشتريات؟ ومن يدير عضويتها؟',
'هل يحتاج المدير أو السكرتارية إلى العمل بالنيابة مع تسجيل اسم المنفذ وصاحب الصلاحية؟'
]),
('4. تسجيل البريد الوارد', [
'ما قنوات وصول الكتب: يدوي، بريد إلكتروني، منصة حكومية، فاكس، واتساب رسمي، مندوب؟',
'من يملك صلاحية تسجيل الوارد؟ وهل التسجيل مركزي في الديوان أم موزع على الأقسام؟',
'ما الصيغة الرسمية لرقم الوارد؟ هل يبدأ من جديد سنويًا؟ وهل توجد سلاسل مختلفة؟',
'ما الحقول الإلزامية: رقم الأصل، تاريخه، الجهة، الموضوع، النوع، السرية، الأولوية، القسم؟',
'هل يسمح بتسجيل كتاب دون ملف مرفق مؤقتًا؟ ومن يتابع استكمال النواقص؟',
'كيف نكتشف الكتاب المكرر؟ وما الإجراء عند اكتشاف تكرار بعد التسجيل؟',
'هل يجب مسح الورق ضوئيًا داخل النظام؟ وما الصيغ والحجم الأقصى وعدد المرفقات؟',
'من يملك تعديل بيانات الكتاب بعد إعطائه رقمًا؟ وهل يحتاج التعديل إلى موافقة؟',
'هل يُطبع باركود أو QR أو ملصق قيد على النسخة الورقية؟ وما البيانات الظاهرة عليه؟'
]),
('5. إنشاء البريد الصادر واعتماده', [
'من يستطيع إنشاء مسودة كتاب صادر؟ وهل توجد قوالب رسمية مختلفة حسب نوع الكتاب؟',
'هل يحرر نص الكتاب داخل النظام أم يُرفع Word/PDF جاهز؟ وأيهما النسخة الرسمية؟',
'ما مسار التدقيق والاعتماد لكل نوع: موظف، رئيس قسم، مدير، رئيس بلدية، شؤون قانونية؟',
'هل تختلف سلسلة الاعتماد حسب القسم أو السرية أو قيمة المعاملة أو الجهة المستلمة؟',
'متى يُحجز رقم الصادر: عند إنشاء المسودة أم بعد الاعتماد النهائي؟ وماذا يحدث للرقم الملغى؟',
'هل يلزم توقيع إلكتروني أو ختم رقمي؟ ومن المخول بالتوقيع؟',
'ما طرق الإرسال التي يجب تسجيلها: مندوب، بريد إلكتروني، منصة، بريد مسجل؟',
'هل نحتاج إثبات استلام ورقم تتبع وتاريخ تسليم واسم المستلم؟',
'كيف تتم إعادة الكتاب للمعدّ للتعديل، وكيف نحفظ الإصدارات والملاحظات دون فقد النسخة السابقة؟'
]),
('6. المراسلات الداخلية والتعاميم', [
'ما أنواع المراسلات الداخلية المعتمدة وتعريف كل نوع؟',
'هل يمكن إرسال المراسلة لموظف، قسم، عدة أقسام، مجموعة ثابتة، أو جميع الموظفين؟',
'هل يحتاج التعميم إلى إثبات قراءة من كل مستلم؟ وماذا يحدث لمن لم يقرأه؟',
'هل التكليف يحتاج تاريخ إنجاز ونتيجة تنفيذ ومرفقات إثبات؟',
'هل يسمح للمستلم بالرد داخل نفس سلسلة المراسلة أم ينشئ مراسلة جديدة مرتبطة؟',
'هل توجد مراسلات داخلية سرية لا تظهر لرئيس القسم أو السكرتارية؟',
'متى تعتبر المراسلة الداخلية مغلقة؟ ومن يملك قرار الإغلاق أو إعادة فتحها؟'
]),
('7. التحويل والتأشير ومسار العمل', [
'إلى من يمكن التحويل: موظف، قسم، مجموعة، لجنة، مدير، أم جهة خارجية؟',
'ما عبارات التأشير الرسمية المطلوبة: للاطلاع، للتنفيذ، للدراسة، للرد، للمتابعة، للعلم؟',
'هل يمكن تحويل الكتاب لأكثر من جهة بالتوازي؟ ومن يملك النسخة الأصلية ومسؤولية الإغلاق؟',
'هل التحويل يكون للعلم فقط أم يتطلب إجراء أو ردًا إلزاميًا؟ وكيف نميز بينهما؟',
'هل يستطيع المستلم إعادة التحويل؟ وهل يحتاج ذلك إلى إذن أو سبب مكتوب؟',
'ما قواعد الإرجاع للمرسل، والاعتذار عن الاستلام، وطلب التمديد، وإعادة الفتح؟',
'هل يمكن تحديد مهلة مختلفة لكل مستلم عند التحويل المتعدد؟',
'من يستطيع تغيير الأولوية أو السرية أو الموعد بعد بدء المعالجة؟',
'ما البيانات التي يجب أن تظهر في سجل الحركة: المحول، المستلم، الإجراء، الوقت، الملاحظة، عنوان IP؟',
'هل نحتاج ربط الكتاب بمراسلات سابقة أو معاملة رئيسية؟ وكيف يتم فك الربط أو تصحيحه؟'
]),
('8. الحالات والمهل والإشعارات', [
'اعتمدوا قائمة الحالات النهائية ومعنى كل حالة ومن ينقل الكتاب إليها.',
'متى يبدأ احتساب المهلة: وقت التسجيل، التحويل، فتح الموظف للكتاب، أم يوم العمل التالي؟',
'هل تُحتسب أيام العطل ونهاية الأسبوع؟ وما التقويم الرسمي المعتمد؟',
'متى يصبح الكتاب متأخرًا؟ وهل توجد فترة سماح حسب الأولوية؟',
'ما قنوات الإشعار المطلوبة: داخل النظام، بريد، SMS، إشعار الموقع الأساسي؟',
'من يستلم تنبيه التأخير والتصعيد، ومتى ينتقل التنبيه لرئيس القسم أو المدير؟',
'هل يستطيع المستخدم إيقاف بعض الإشعارات؟ وما الإشعارات الإلزامية التي لا يمكن تعطيلها؟',
'هل نحتاج تذكيرًا يدويًا وأجندة مهام يومية وأسبوعية للموظف؟'
]),
('9. السرية والصلاحيات والتدقيق', [
'ما درجات السرية المعتمدة رسميًا، وما الفرق العملي في الصلاحيات بين كل درجة؟',
'هل السرية تُطبق على بيانات الكتاب فقط أم على الملف والملاحظات وسجل الحركة أيضًا؟',
'من يستطيع تنزيل أو طباعة أو مشاركة الملفات السرية؟ وهل نحتاج علامة مائية باسم المستخدم؟',
'هل توجد صلاحيات منفصلة للعرض والإنشاء والتعديل والتحويل والاعتماد والحذف والأرشفة والتصدير؟',
'هل يسمح بالحذف الحقيقي؟ أم نستخدم إلغاء/سحب مع بقاء السجل لأغراض التدقيق؟',
'كم مدة الجلسة؟ وهل نحتاج تحققًا إضافيًا عند الاعتماد أو فتح كتاب سري؟',
'ما العمليات التي يجب تسجيلها في سجل غير قابل للتعديل؟ وكم مدة الاحتفاظ به؟',
'من يراجع سجل العمليات؟ وهل توجد تقارير لمحاولات الدخول والفشل والتصرفات غير المعتادة؟'
]),
('10. الأرشيف وإدارة الوثائق', [
'ما خطة تصنيف الأرشيف: سنة، نوع، جهة، قسم، موضوع، ملف، مشروع، لجنة؟',
'هل الأرشفة تلقائية بعد الإغلاق أم تحتاج اعتماد مسؤول الأرشيف؟',
'ما قواعد تسمية الملفات وصيغها المقبولة والحجم الأقصى وعدد المرفقات؟',
'هل نحتاج OCR للبحث داخل الملفات الممسوحة؟ وهل الوثائق العربية القديمة واضحة بما يكفي؟',
'هل نحتاج إدارة نسخ وإصدارات ومنع استبدال الملف النهائي دون سجل؟',
'ما مدد الاحتفاظ بكل نوع وثيقة؟ وما آلية الإتلاف القانوني بعد انتهاء المدة؟',
'هل توجد ملفات لا يجوز تنزيلها أو يجب حفظها في خادم منفصل؟',
'كيف نستعيد كتابًا مؤرشفًا أو نعيد فتحه؟ ومن يعتمد العملية؟',
'هل نحتاج تصدير ملف معاملة كامل يشمل البيانات والمرفقات وسجل الحركة؟'
]),
('11. البحث والتقارير ولوحات المتابعة', [
'بأي حقول يبحث الموظفون غالبًا: الرقم، الموضوع، الجهة، الشخص، التاريخ، النص داخل الملف؟',
'هل يجب أن تظهر نتائج البحث فقط ضمن صلاحيات المستخدم، حتى لو عرف رقم الكتاب؟',
'ما التقارير اليومية والشهرية المطلوبة للديوان والإدارة والأقسام؟',
'كيف تقيسون الإنجاز: عدد الكتب، نسبة الإغلاق ضمن المهلة، متوسط زمن المعالجة، المتأخر؟',
'هل نحتاج تقرير حمل العمل لكل موظف وقسم؟ ومن يحق له رؤيته؟',
'ما صيغ التصدير المطلوبة: PDF، Excel، طباعة رسمية؟ وهل يلزم شعار وختم وتوقيع؟',
'ما المؤشرات التي يجب أن تظهر في لوحة المدير، ولوحة رئيس القسم، ولوحة الموظف؟',
'هل توجد تقارير رقابية أو قانونية بصيغة محددة يجب مطابقتها؟'
]),
('12. التكامل والتشغيل والاستمرارية', [
'ما التقنية المستخدمة في موقع البلدية الأساسي؟ ومن المسؤول الفني عن دمج الموديول؟',
'كيف سيستلم الموديول هوية المستخدم وقسمه وصلاحياته من النظام الأساسي؟',
'ما الأنظمة المطلوب الربط معها: المستخدمون، البريد، الرسائل، الموارد البشرية، منصة حكومية؟',
'أين ستُحفظ قاعدة البيانات والملفات؟ داخل البلدية أم استضافة خارجية؟',
'هل يوجد خادم تجريبي وخادم إنتاج؟ ومن يملك صلاحية النشر والتحديث؟',
'ما سياسة النسخ الاحتياطي: التكرار، المكان، التشفير، ومدة الاحتفاظ؟ وهل تم اختبار الاستعادة؟',
'ما المطلوب عند انقطاع الإنترنت أو الكهرباء؟ وهل يلزم عمل محلي مؤقت؟',
'كم عدد المستخدمين المتزامنين وما مستوى الأداء المقبول لفتح الصفحات والملفات والبحث؟',
'من يقدم الدعم بعد الإطلاق؟ وما تصنيف الأعطال ووقت الاستجابة المقبول؟',
'ما خطة تدريب الديوان والمديرين والموظفين؟ ومن يعتمد دليل الاستخدام؟'
]),
('13. الاعتماد والأولويات قبل البرمجة', [
'من سيوقّع وثيقة المتطلبات النهائية ممثلًا عن البلدية؟',
'ما سيناريوهات القبول التي يجب تنفيذها أمام المسؤول قبل الاستلام؟',
'ما البيانات التجريبية المسموح استخدامها دون كشف مراسلات حقيقية أو معلومات حساسة؟',
'ما الأولوية لكل وحدة: ضروري للإطلاق، مهم، تحسين لاحق؟',
'ما الموعد المستهدف للإطلاق، وما القيود التي قد تؤثر عليه؟',
'كيف تُطلب التغييرات بعد الاعتماد؟ ومن يوافق على أثرها في الوقت والتكلفة؟',
'هل توجد موافقات قانونية أو سياسة خصوصية أو تعليمات حكومية يجب مراجعتها قبل التشغيل؟'
])]

doc=Document()
sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11)
sec.top_margin=Inches(.78); sec.bottom_margin=Inches(.72); sec.left_margin=Inches(.82); sec.right_margin=Inches(.82)
sec.header_distance=Inches(.3); sec.footer_distance=Inches(.3)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(NAVY)
normal._element.rPr.rFonts.set(qn('w:cs'),'Arial'); normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial')
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for name,size,color,before,after in [('Heading 1',16,BLUE,18,10),('Heading 2',13,BLUE,14,7),('Heading 3',12,'1F4D78',10,5)]:
    st=styles[name]; st.font.name='Calibri'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
    st._element.rPr.rFonts.set(qn('w:cs'),'Arial'); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
question_style=styles.add_style('Question',WD_STYLE_TYPE.PARAGRAPH); question_style.font.name='Arial'; question_style.font.size=Pt(11); question_style.font.bold=True; question_style.font.color.rgb=RGBColor.from_string(NAVY); question_style.paragraph_format.space_before=Pt(8); question_style.paragraph_format.space_after=Pt(3); question_style.paragraph_format.keep_with_next=True
answer_style=styles.add_style('Answer Label',WD_STYLE_TYPE.PARAGRAPH); answer_style.font.name='Arial'; answer_style.font.size=Pt(9); answer_style.font.color.rgb=RGBColor.from_string(MUTED); answer_style.paragraph_format.space_after=Pt(2); answer_style.paragraph_format.keep_with_next=True

def rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p.alignment=align
    pPr=p._p.get_or_add_pPr(); bidi=pPr.find(qn('w:bidi'))
    if bidi is None: pPr.append(OxmlElement('w:bidi'))
    return p
def set_cell_shading(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)
def set_cell_margins(cell,top=100,start=140,bottom=100,end=140):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); mar=tcPr.first_child_found_in('w:tcMar')
    if mar is None: mar=OxmlElement('w:tcMar'); tcPr.append(mar)
    for edge,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=mar.find(qn('w:'+edge))
        if node is None: node=OxmlElement('w:'+edge); mar.append(node)
        node.set(qn('w:w'),str(val)); node.set(qn('w:type'),'dxa')
def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)
def add_page_field(p):
    r=p.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); r._r.append(fld)
def add_answer_box(lines=2):
    table=doc.add_table(rows=1,cols=1); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False; table.columns[0].width=Inches(6.5)
    cell=table.cell(0,0); cell.width=Inches(6.5); set_cell_shading(cell,'FAFBFC'); set_cell_margins(cell,100,140,100,140)
    cell.text='الإجابة / القرار:'
    for p in cell.paragraphs: rtl(p); p.style='Answer Label'
    for _ in range(lines):
        p=rtl(cell.add_paragraph(' ')); p.paragraph_format.space_after=Pt(2)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

header=sec.header
p=rtl(header.paragraphs[0]); p.text='بلدية المغازي | مشروع نظام المراسلات والأرشيف الإلكتروني'; p.style=styles['Normal']; p.runs[0].font.size=Pt(8.5); p.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
footer=sec.footer
fp=footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=fp.add_run('وثيقة تحليل متطلبات - صفحة '); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(MUTED); add_page_field(fp)

if LOGO.exists():
    p=rtl(doc.add_paragraph(),WD_ALIGN_PARAGRAPH.CENTER); picture=p.add_run().add_picture(str(LOGO),width=Inches(.85)); picture._inline.docPr.set('descr','شعار بلدية المغازي'); picture._inline.docPr.set('title','شعار بلدية المغازي')
p=rtl(doc.add_paragraph(),WD_ALIGN_PARAGRAPH.CENTER); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
r=p.add_run('دليل مقابلة تحليل المتطلبات'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(25); r.font.color.rgb=RGBColor.from_string(MAROON)
p=rtl(doc.add_paragraph(),WD_ALIGN_PARAGRAPH.CENTER); p.paragraph_format.space_after=Pt(5)
r=p.add_run('نظام المراسلات والأرشيف الإلكتروني - بلدية المغازي'); r.bold=True; r.font.name='Arial'; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(NAVY)
p=rtl(doc.add_paragraph(),WD_ALIGN_PARAGRAPH.CENTER); p.paragraph_format.space_after=Pt(18)
r=p.add_run('أسئلة اعتماد دورة العمل، الصلاحيات، الأرشفة، التكامل ومعايير القبول'); r.font.name='Arial'; r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string(MUTED)

meta=doc.add_table(rows=4,cols=2); meta.alignment=WD_TABLE_ALIGNMENT.CENTER; meta.autofit=False
for row in meta.rows:
    row.cells[0].width=Inches(1.55); row.cells[1].width=Inches(4.95)
for i,(label,value) in enumerate([('تاريخ الاجتماع',''),('اسم المسؤول وصفته',''),('الحضور',''),('إصدار الوثيقة','1.0 - للتحليل والاعتماد')]):
    meta.cell(i,0).text=label; meta.cell(i,1).text=value
    for c in meta.rows[i].cells:
        set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in c.paragraphs: rtl(p)
    set_cell_shading(meta.cell(i,0),PALE); meta.cell(i,0).paragraphs[0].runs[0].bold=True; meta.cell(i,0).paragraphs[0].runs[0].font.color.rgb=RGBColor.from_string(MAROON)

doc.add_paragraph()
call=doc.add_table(rows=1,cols=1); call.alignment=WD_TABLE_ALIGNMENT.CENTER; call.autofit=False; call.columns[0].width=Inches(6.5)
cell=call.cell(0,0); set_cell_shading(cell,'FFF8E8'); set_cell_margins(cell,180,180,180,180)
p=rtl(cell.paragraphs[0]); p.paragraph_format.space_after=Pt(3); r=p.add_run('قاعدة الجلسة'); r.bold=True; r.font.color.rgb=RGBColor.from_string(GOLD)
p=rtl(cell.add_paragraph('لا نكتفي بإجابة «نعم». اطلبوا مثالًا واقعيًا، اسم صاحب الصلاحية، الاستثناءات، والقرار النهائي لكل سؤال. أي نقطة غير محسومة تُسجل كقرار معلّق مع مسؤول وموعد حسم.'))

doc.add_paragraph()
agenda=doc.add_table(rows=2,cols=4); agenda.alignment=WD_TABLE_ALIGNMENT.CENTER; agenda.autofit=False
for c,w in zip(agenda.columns,[1.625]*4): c.width=Inches(w)
for idx,(num,label) in enumerate([('01','فهم الواقع'),('02','اعتماد المسارات'),('03','الأمن والتكامل'),('04','الأولويات والقبول')]):
    c=agenda.cell(0,idx); c.text=num; set_cell_shading(c,MAROON)
    for p in c.paragraphs: p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].font.color.rgb=RGBColor.from_string(WHITE); p.runs[0].bold=True
    c2=agenda.cell(1,idx); c2.text=label; set_cell_shading(c2,LIGHT)
    for p in c2.paragraphs: rtl(p,WD_ALIGN_PARAGRAPH.CENTER); p.runs[0].font.size=Pt(9); p.runs[0].bold=True

doc.add_page_break()
p=rtl(doc.add_paragraph('طريقة استخدام الدليل',style='Heading 1'))
for txt in ['اعقدوا جلسة مع الديوان والإدارة ورؤساء الأقسام والأرشيف وتقنية المعلومات، ولا تعتمدوا على شخص واحد فقط.','اطلبوا وثائق وأمثلة حقيقية: كتاب وارد، كتاب صادر، تعميم، سجل يدوي، تقرير، وقالب رسمي.','دوّنوا القرار المعتمد لا الاقتراح الأول، وحددوا صاحب القرار والموعد لأي نقطة معلقة.','بعد الجلسات، حوّلوا الإجابات إلى وثيقة متطلبات ومسارات وصلاحيات ونماذج قبول يوقع عليها المسؤول.']:
    p=rtl(doc.add_paragraph(style='List Bullet')); p.add_run(txt)

qnum=1
for title,questions in sections:
    rtl(doc.add_paragraph(title,style='Heading 1'))
    intro=rtl(doc.add_paragraph('سجّل الإجابة المعتمدة، الاستثناءات، واسم صاحب القرار.')); intro.runs[0].italic=True; intro.runs[0].font.color.rgb=RGBColor.from_string(MUTED)
    for question in questions:
        p=rtl(doc.add_paragraph(style='Question')); p.add_run(f'{qnum}. {question}')
        add_answer_box(1)
        qnum+=1

doc.add_page_break()
rtl(doc.add_paragraph('سجل القرارات المعلّقة',style='Heading 1'))
t=doc.add_table(rows=1,cols=5); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
widths=[.55,2.55,1.15,1.15,1.1]
for c,w in zip(t.columns,widths): c.width=Inches(w)
for cell,text in zip(t.rows[0].cells,['#','القرار المعلّق','المسؤول','موعد الحسم','الحالة']):
    cell.text=text; set_cell_shading(cell,MAROON); set_cell_margins(cell)
    for p in cell.paragraphs: rtl(p,WD_ALIGN_PARAGRAPH.CENTER); p.runs[0].font.color.rgb=RGBColor.from_string(WHITE); p.runs[0].bold=True
set_repeat_table_header(t.rows[0])
for i in range(1,13):
    cells=t.add_row().cells
    vals=[str(i),'','','','معلّق']
    for cell,val,w in zip(cells,vals,widths):
        cell.text=val; cell.width=Inches(w); set_cell_margins(cell,160,100,160,100)
        for p in cell.paragraphs: rtl(p,WD_ALIGN_PARAGRAPH.CENTER if cell in (cells[0],cells[4]) else WD_ALIGN_PARAGRAPH.RIGHT)

rtl(doc.add_paragraph('قائمة الاعتماد قبل بدء التطوير',style='Heading 1'))
checks=['النطاق والمرحلة الأولى معتمدان.','الهيكل والأدوار ومصدر المستخدمين معتمدة.','حقول الوارد والصادر والترقيم معتمدة.','مسارات التحويل والاعتماد والاستثناءات معتمدة.','الحالات والمهل والتصعيدات معتمدة.','درجات السرية ومصفوفة الصلاحيات معتمدة.','خطة الأرشيف والاستيراد والاحتفاظ معتمدة.','التقارير ولوحات المتابعة معتمدة.','واجهات الربط والبنية والاستضافة والنسخ الاحتياطي معتمدة.','سيناريوهات القبول وخطة التدريب والإطلاق معتمدة.']
for text in checks:
    p=rtl(doc.add_paragraph(style='List Bullet')); p.add_run('☐ '+text)

rtl(doc.add_paragraph('اعتماد المسؤول',style='Heading 1'))
sign=doc.add_table(rows=3,cols=2); sign.alignment=WD_TABLE_ALIGNMENT.CENTER; sign.autofit=False
for row in sign.rows:
    row.cells[0].width=Inches(1.6); row.cells[1].width=Inches(4.9)
for i,(label,value) in enumerate([('الاسم والصفة',''),('التوقيع',''),('التاريخ','')]):
    sign.cell(i,0).text=label; sign.cell(i,1).text=value
    set_cell_shading(sign.cell(i,0),PALE)
    for c in sign.rows[i].cells:
        set_cell_margins(c,180,140,180,140)
        for p in c.paragraphs: rtl(p)

core=doc.core_properties; core.title='دليل مقابلة تحليل متطلبات نظام المراسلات والأرشيف الإلكتروني'; core.subject='بلدية المغازي'; core.author='بلدية المغازي'; core.keywords='مراسلات، أرشيف، تحليل متطلبات، بلدية المغازي'
doc.save(OUT)
print('DOCX_CREATED')
